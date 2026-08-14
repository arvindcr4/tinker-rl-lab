from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(
    "/Users/arvind/Developer/tinker-rl-lab/outputs/"
    "NeurIPS_12_Ideas_E1_E14_Read_Aloud_Transcript_2026-08-09.docx"
)

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(28, 34, 44)
MUTED = RGBColor(96, 105, 118)
LIGHT = "E8EEF5"


SLIDES = [
    (
        "Slide 1 - How the review reshaped the research",
        [
            "Hello. Today I want to explain how the NeurIPS review changed my research program. This is not a story about defending every old claim. It is a story about making the work easier to understand, easier to test, and harder to overstate.",
            "I will cover four things. First, what the reviewers liked and what worried them. Second, how I corrected twelve research ideas. Third, why only two directions now deserve serious paper-level work. And fourth, how the new E1 through E14 evaluation plan tests whether the models are useful beyond one math benchmark.",
        ],
    ),
    (
        "Slide 2 - What the NeurIPS reviewers said",
        [
            "The reviewers did see value in submission number 36320. They agreed that reinforcement-learning training curves can be misleading. They also saw value in Z-V-F and G-U because these diagnostics are cheap and easy to monitor. One reviewer specifically recognized the importance of separating online reward, held-out capability, and algorithm labels.",
            "But the reviewers could not trust the overall story. The writing was hard to follow. The abstract carried too many equations and details. The runner was not standard G-R-P-O, but the paper did not make that boundary clear enough. Results from different models, tasks, and software stacks were placed too close together. Some rows had weak or conflicting provenance. And the paper called itself use-inspired without showing a completed real-world outcome.",
            "So the central idea was interesting, but the evidence was too fragmented for the broad narrative.",
        ],
    ),
    (
        "Slide 3 - I treated the review like a bug report",
        [
            "I responded in four steps: admit, audit, narrow, and rebuild.",
            "First, I admitted which claims were unsupported. Second, I traced important numbers back to the exact run, seed, model, evaluator, and artifact. Third, I kept only the claims that the evidence could actually carry. Finally, I designed stronger tests for future claims.",
            "This meant withdrawing several attractive statements. I removed the pooled mean over heterogeneous checkpoints. I quarantined the conflicting P-P-O row. I withdrew an unsupported five-seed transfer comparison. I stopped treating high Z-V-F as proof of high-reward saturation. And I removed broad deployment or usefulness claims that had not been prospectively tested.",
            "The biggest correction was simple: a clear statement that I do not know yet is stronger than a shaky positive result.",
        ],
    ),
    (
        "Slide 4 - Corrections to research ideas P1 through P6",
        [
            "The first six research ideas now have smaller and clearer jobs.",
            "P1, scaling, no longer pools selected checkpoints or claims a universal scaling law. It is now a limits and identifiability audit. P2 treats Z-V-F as a descriptive diagnostic, not the whole gradient and not a universal predictor. P3 reports group-size results only inside the measured token budget; it makes no one-size-fits-all recommendation. P4 reports a bounded null result under a two-hundred-token cap, not a general theory of length bias.",
            "P5 is the minimum reporting idea. It requires researchers to state provenance, the quantity being estimated, missing experiment cells, and held-out pass-at-k results. P6 is the run registry. When two records conflict, the registry quarantines them instead of averaging the disagreement away.",
        ],
    ),
    (
        "Slide 5 - Corrections to research ideas P7 through P12",
        [
            "The next six ideas were narrowed in the same way.",
            "P7, the Z-V-F controller, is now a retrospective audit and a future test plan. It does not claim that the controller improves performance. P8 remains an exploratory workshop artifact, with no broad rankings. P9 uses evidence tiers so uncertain runs cannot silently support strong claims. P10 limits its proofs to centered reward contrast; it does not claim to describe every possible gradient term.",
            "P11 reports a completed forty-unit, single-stack audit, but the algorithm comparisons remain inconclusive. P12 introduces signal-starvation measurements and the triage R-L routing proposal, but it does not claim that P-P-O or S-A-O training improved. Those outcomes still need prospective experiments.",
            "These twelve items are best described as research ideas or manuscript directions. They are not twelve equally strong papers.",
        ],
    ),
    (
        "Slide 6 - The two serious paper directions",
        [
            "After the audit, two directions remain strong enough for serious paper work.",
            "Paper direction A is treatment verification. Its simple message is that an algorithm name is only a label. Before comparing results, we must prove what code actually ran and bind the claim to the exact stack, evaluator, seed structure, and evidence. This direction combines the reporting standard, the registry, and the single-stack survival audit. It is defensible now as a methodology, artifact, or reproducibility paper.",
            "Paper direction B is triage R-L. Its idea is that weak learning signals can have different causes. An example may already be solved. It may be failing. A critic may be unreliable. A clipping rule may discard useful signal. Or the trajectory may be unsafe or invalid. These cases should not receive the same action.",
            "Triage R-L is the higher-upside flagship direction, but it is gated. It becomes flagship-worthy only if a preregistered, matched-budget experiment beats static and simpler baselines. Today it is worthy of testing, not yet worthy of a success claim.",
        ],
    ),
    (
        "Slide 7 - Why GSM8K alone was not enough",
        [
            "The next correction was about usefulness. G-S-M-eight-K is like one math worksheet. Passing it can show a narrow reasoning skill, but it cannot show that a model can do the work required by many real companies.",
            "Real work includes using tools, editing repositories, browsing, changing environment state, handling finance tasks, building artifacts, writing chip-design code, staying safe, and completing long sequences of actions.",
            "So G-S-M-eight-K is now calibration only. The new rule is to train across several task families, evaluate on unseen families, and never hide a weak domain inside one overall average. E1 through E14 turns the word useful from a slogan into fourteen separate report cards.",
        ],
    ),
    (
        "Slide 8 - E1 through E7",
        [
            "The first seven evaluation lanes cover code, enterprise work, browsers, finance, and security.",
            "E1 is S-W-E-bench Pro for difficult repository repair. E2 is Frontier S-W-E. E3 is S-D-A-B for production-system and enterprise work. E4 is Banker Tool Bench for finance tools and state changes. E5 is APEX Agents for long professional workflows. E6 is WebBench for browser and computer use. E7 is BinaryAudit for security and binary analysis.",
            "Most of these lanes are still blocked before a real model score. E5 and E7 have partial harness evidence. Partial means that plumbing or a verifier worked on a controlled fixture. It does not mean the model passed the benchmark. Blocked also does not mean no work was done; it means the exact evidence package required for a score is not complete.",
        ],
    ),
    (
        "Slide 9 - E8 through E14",
        [
            "The second group expands into science, machine learning, safety, chip design, visual applications, games, and hard mathematics.",
            "E8 is LifeSciBench. E9 is M-L-E-bench. E10 is AgentHarm. E11 is VerilogEval for chip-design code. E12 is AppBench. E13 is OpenReward Games. E14 is FrontierMath.",
            "At present, E11 is the only lane with a trained-model score, and even that result is deliberately small. E9 and E10 have partial harness evidence. The remaining lanes are blocked on exact datasets, permissions, environments, private splits, or native verifiers.",
            "Again, the status labels protect the claim boundary. A public sample, mock model, synthetic fixture, or reference answer cannot be promoted into a benchmark score.",
        ],
    ),
    (
        "Slide 10 - What has actually run",
        [
            "There is one small but real portfolio result. I trained a Qwen three-point-six, thirty-five-billion mixture-of-experts model for forty reinforcement-learning steps on API tool tasks and software-repair examples. The same thirty-two unseen examples were then evaluated with the same sampling and scoring setup.",
            "The starting model had a mean reward of zero-point-three-nine-six. The trained model reached zero-point-five-nine-nine. That is an absolute increase of about zero-point-two-zero-three, or about fifty-one percent relative improvement.",
            "This is encouraging, but it is not an official S-W-E-bench Pro result. The software-repair component uses a dense partial-credit checker. It asks whether the output looks like a valid patch, touches the right files, and resembles the reference. That is useful local evidence, but it is not the full benchmark verifier.",
            "For E11, the trained model passed four out of four sampled Verilog tasks, with zero extraction failures. The full suite contains three hundred and twelve prompts, so four out of four is a smoke test, not a full VerilogEval score.",
            "The honest summary is: one of fourteen evaluation lanes has a trained-model score today, and thirteen still need exact model-level results.",
        ],
    ),
    (
        "Slide 11 - Closing",
        [
            "The new research program can be summarized in four words: less claiming, more proving.",
            "The twelve research ideas now say more clearly what their evidence supports. Two paper directions carry the strongest novelty. Fourteen evaluation suites test whether the work matters beyond one short math benchmark.",
            "Before the next strong claim, I require exact data, an unseen split, a native verifier, a Weights and Biases run, a Hugging Face checkpoint, and a matched baseline. These are not paperwork. They are the chain that lets another person check the result.",
            "The goal is not to show fourteen green boxes today. The goal is to produce fourteen results that I can defend tomorrow.",
            "Thank you.",
        ],
    ),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def set_font(run, size: float, color: RGBColor = INK, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h1.font.size = Pt(16)
    h1.font.color.rgb = BLUE
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h2.font.size = Pt(13)
    h2.font.color.rgb = BLUE
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)


def add_script_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.widow_control = True
    run = p.add_run(text)
    set_font(run, 11.5)


def build() -> None:
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header_p = section.header.paragraphs[0]
    header_p.text = "NEURIPS REVIEW RESEARCH UPDATE  |  READ-ALOUD TRANSCRIPT"
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header_p.runs:
        set_font(run, 8.5, MUTED, bold=True)
    add_page_number(section.footer.paragraphs[0])

    for _ in range(5):
        doc.add_paragraph()

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("PLAIN-LANGUAGE SPEAKING SCRIPT")
    set_font(run, 10, BLUE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("NeurIPS Review, 12 Research Ideas,\nand the E1-E14 Expansion")
    set_font(run, 28, DARK_BLUE, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(26)
    run = subtitle.add_run("A transcript aligned to the 11-slide dark presentation")
    set_font(run, 14, MUTED, italic=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(5)
    run = meta.add_run("Approximate speaking time: 9-11 minutes")
    set_font(run, 11, INK, bold=True)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run("Read the main paragraphs aloud. Gray bracketed notes are optional cues.")
    set_font(run, 10, MUTED)

    doc.add_page_break()

    doc.add_heading("Quick pronunciation guide - do not read aloud", level=1)
    terms = [
        ("ZVF", 'say "Z-V-F"'),
        ("GU", 'say "G-U"'),
        ("GRPO", 'say "G-R-P-O"'),
        ("TRIAGE-RL", 'say "triage R-L"'),
        ("W&B", 'say "Weights and Biases"'),
        ("HF", 'say "Hugging Face"'),
        ("E1-E14", 'say "E one through E fourteen"'),
    ]
    table = doc.add_table(rows=len(terms), cols=2)
    table.autofit = False
    for i, (term, speech) in enumerate(terms):
        table.rows[i].cells[0].width = Inches(1.35)
        table.rows[i].cells[1].width = Inches(5.15)
        table.rows[i].cells[0].text = term
        table.rows[i].cells[1].text = speech
        set_cell_shading(table.rows[i].cells[0], LIGHT)
        for j, cell in enumerate(table.rows[i].cells):
            cell.vertical_alignment = 1
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    set_font(run, 10.5, INK, bold=(j == 0))

    cue = doc.add_paragraph()
    cue.paragraph_format.space_before = Pt(12)
    cue.paragraph_format.space_after = Pt(6)
    run = cue.add_run("[Optional cue: pause briefly when moving to a new slide.]")
    set_font(run, 10, MUTED, italic=True)

    for title_text, paragraphs in SLIDES:
        doc.add_heading(title_text, level=1)
        for paragraph in paragraphs:
            add_script_paragraph(doc, paragraph)

    doc.add_page_break()
    doc.add_heading("Source basis - do not read aloud", level=1)
    sources = [
        "/Users/arvind/.codex/attachments/191be32b-d5ed-4464-a778-7ed4c495f73f/pasted-text.txt",
        "/Users/arvind/Developer/tinker-rl-lab/platform_hybrid/paper/PAPERS_README.md",
        "/Users/arvind/Developer/tinker-rl-lab/platform_hybrid/paper/REVIEWER_36320_CORRECTION_MANIFEST.md",
        "/Users/arvind/Developer/tinker-rl-lab/BREAKTHROUGH_CHASE_18_ARTIFACTS.md",
        "/Users/arvind/Developer/tinker-rl-lab/zvf-program/flagship/PAVLOVS_LIST_TASK_CONTRACT.md",
        "/Users/arvind/Developer/tinker-rl-lab/outputs/pavlov_portfolio_eval/base_reasoning_stripped_seed1810.json",
        "/Users/arvind/Developer/tinker-rl-lab/outputs/pavlov_portfolio_eval/trained_step40_seed1810.json",
        "/Users/arvind/Developer/tinker-rl-lab/outputs/e11_verilog_eval/e11_trained_step40_receipt.json",
    ]
    for source in sources:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(source)
        set_font(run, 8.5, MUTED)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
